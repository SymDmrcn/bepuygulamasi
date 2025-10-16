import json
import streamlit as st
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document

firebase_config = st.secrets["firebase_config"]
cred = credentials.Certificate(dict(firebase_config))
if not firebase_admin._apps:
    firebase_admin.initialize_app(cred)
db = firestore.client()

# Cache'e TTL ekle (5 saniyede bir yenilenir - geliÅŸtirme iÃ§in)
@st.cache_data(ttl=5)
def verileri_cek():
    hedefler_ref = db.collection('hedefler')
    docs = hedefler_ref.stream()
    grouped_data = {}
    
    for doc in docs:
        data = doc.to_dict()
        grup = data.get("grup", "").strip()
        ders = data.get("ders", "").strip()
        
        if not grup or not ders:
            continue
            
        if grup not in grouped_data:
            grouped_data[grup] = {}
        
        # KÄ±sa vadeli hedefleri birleÅŸtir
        kisa_vadeli = []
        kisa_vadeli.extend(data.get("kisa_vadeli_hedefler", []))
        kisa_vadeli.extend(data.get("kÄ±sa_vadeli_hedefler", []))
        
        # Uzun vadeli hedefleri birleÅŸtir
        uzun_vadeli = []
        uzun_vadeli.extend(data.get("uzun_vadeli_hedefler", []))
        uzun_vadeli.extend(data.get("uzak_vadeli_hedefler", []))
        
        # Ã–ÄŸretimsel hedefleri birleÅŸtir  
        ogretimsel = []
        ogretimsel.extend(data.get("ogretimsel_hedefler", []))
        ogretimsel.extend(data.get("Ã¶ÄŸretimsel_hedefler", []))
        
        grouped_data[grup][ders] = {
            "KISA_VADELI_HEDEFLER": list(set(kisa_vadeli)),  # Dublicateleri kaldÄ±r
            "UZUN_VADELI_HEDEFLER": list(set(uzun_vadeli)),  # Dublicateleri kaldÄ±r
            "OGRETIMSEL_HEDEFLER": list(set(ogretimsel))  # Dublicateleri kaldÄ±r
        }
    
    return grouped_data

# ğŸ¨ ArayÃ¼z
st.set_page_config(page_title="BEP OluÅŸturucu", layout="centered")
st.title("ğŸ“˜ TUZLA BÄ°LSEM BireyselleÅŸtirilmiÅŸ EÄŸitim PlanÄ± Otomasyonu (BEP)")

# Cache temizleme butonu
if st.sidebar.button("ğŸ”„ Verileri Yenile"):
    st.cache_data.clear()
    st.rerun()

# Debug bilgileri (geliÅŸtirme aÅŸamasÄ±nda)
if st.sidebar.checkbox("ğŸ” Debug Modu"):
    st.sidebar.write("**Cache Durumu:**")
    st.sidebar.write(f"Cache TTL: 30 saniye")
    
grouped_data = verileri_cek()

# Debug modu aktifse verileri gÃ¶ster
if st.sidebar.checkbox("ğŸ” Debug Modu"):
    st.sidebar.write("**Bulunan Gruplar:**")
    for grup in grouped_data.keys():
        st.sidebar.write(f"- {grup}")
        for ders in grouped_data[grup].keys():
            st.sidebar.write(f"  â””â”€ {ders}")

teacher = st.text_input("ğŸ‘©â€ğŸ« Ã–ÄŸretmen AdÄ±")
student = st.text_input("ğŸ‘§ Ã–ÄŸrenci AdÄ±")

if grouped_data:
    group = st.selectbox("ğŸ¯ Grup SeÃ§in", list(grouped_data.keys()))
    lesson = st.selectbox("ğŸ“š Ders SeÃ§in", list(grouped_data[group].keys()))
    
    hedefler = grouped_data[group][lesson]
    
    # DÃ¼zeltilmiÅŸ multiselect'ler
    short_selected = st.multiselect("ğŸ“ KÄ±sa Vadeli Hedefler", hedefler["KISA_VADELI_HEDEFLER"])
    long_selected = st.multiselect("ğŸ§­ Uzun Vadeli Hedefler", hedefler["UZUN_VADELI_HEDEFLER"])
    teach_selected = st.multiselect("ğŸ“– Ã–ÄŸretimsel Hedefler", hedefler["OGRETIMSEL_HEDEFLER"])
    
    if st.button("ğŸ“„ Word Belgesi OluÅŸtur"):
        if not teacher or not student:
            st.error("LÃ¼tfen Ã¶ÄŸretmen ve Ã¶ÄŸrenci adlarÄ±nÄ± giriniz.")
        else:
            doc = Document()
            doc.add_heading("BireyselleÅŸtirilmiÅŸ EÄŸitim PlanÄ±", level=1)
            doc.add_paragraph(f"Ã–ÄŸretmen: {teacher}")
            doc.add_paragraph(f"Ã–ÄŸrenci: {student}")
            doc.add_paragraph(f"Grup: {group}")
            doc.add_paragraph(f"Ders: {lesson}")
            
            if short_selected:
                doc.add_heading("KÄ±sa Vadeli Hedefler", level=2)
                for item in short_selected:
                    doc.add_paragraph(f"- {item}")
                    
            if long_selected:
                doc.add_heading("Uzun Vadeli Hedefler", level=2)
                for item in long_selected:
                    doc.add_paragraph(f"- {item}")
                    
            if teach_selected:
                doc.add_heading("Ã–ÄŸretimsel Hedefler", level=2)
                for item in teach_selected:
                    doc.add_paragraph(f"- {item}")
            
            file_name = f"{student.replace(' ', '_')}_bep.docx"
            doc.save(file_name)
            
            with open(file_name, "rb") as f:
                st.download_button("ğŸ“¥ Word Belgesini Ä°ndir", f, file_name=file_name)
else:
    st.warning("HenÃ¼z Firestore'dan veri alÄ±namadÄ±.")
